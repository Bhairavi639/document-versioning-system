import os
import time
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import pandas as pd

REPORTS_FOLDER = "reports"

def add_highlighted_text(paragraph, text):
    words = text.split()

    for word in words:
        run = paragraph.add_run(word + " ")

        # Added word
        if word.startswith("[") and word.endswith("]"):
            clean_word = word[1:-1]
            run.text = clean_word + " "
            run.font.color.rgb = RGBColor(0, 128, 0)  # green

        else:
            run.text = word + " "
            
def generate_docx_report(diff_data, file1_name, file2_name):
    os.makedirs(REPORTS_FOLDER, exist_ok=True)

    filename = f"comparison_report_{int(time.time())}.docx"
    file_path = os.path.join(REPORTS_FOLDER, filename)

    doc = Document()

    doc.add_heading("Document Comparison Report", 0)

    doc.add_paragraph(f"File 1: {file1_name}")
    doc.add_paragraph(f"File 2: {file2_name}")

    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    headers = table.rows[0].cells

    headers[0].text = "Sheet"
    headers[1].text = "Key"
    headers[2].text = "Column"
    headers[3].text = "Change Type"
    headers[4].text = "Old Value"
    headers[5].text = "New Value"

    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER



    for change in diff_data:

        row = table.add_row().cells

        row[0].text = str(change.get("sheet", ""))
        row[1].text = str(change.get("key", ""))
        row[2].text = str(change.get("column", ""))
        row[3].text = str(change.get("type", ""))

        old_para = row[4].paragraphs[0]
        new_para = row[5].paragraphs[0]

        add_highlighted_text(
            old_para,
            str(change.get("old", ""))
        )

        add_highlighted_text(
            new_para,
            str(change.get("new", ""))
        )

    doc.save(file_path)

    return file_path

def generate_csv_report(diff_data):

    os.makedirs(REPORTS_FOLDER, exist_ok=True)

    filename = f"comparison_report_{int(time.time())}.csv"

    file_path = os.path.join(
        REPORTS_FOLDER,
        filename
    )

    df = pd.DataFrame(diff_data)

    df.to_csv(
        file_path,
        index=False
    )

    return file_path