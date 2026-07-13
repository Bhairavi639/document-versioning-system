from docx import Document


def read_docx(file_path):
    """
    Read a DOCX file and return a list of non-empty paragraphs.
    """

    doc = Document(file_path)

    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()

        if text:
            paragraphs.append(text)

    # ✅ NEW: Read tables
    for table in doc.tables:
        for row in table.rows:
            row_data = []

            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_data.append(cell_text)

            if row_data:
                # Convert row into a line
                paragraphs.append(" | ".join(row_data))

    return paragraphs   # ✅ ADD THIS